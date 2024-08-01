"""
 A Utility to check and provide a printable list of proposed and ready tickets ready for prioritisation.
 To use this the python-docx package needs to be installed.
 There is limited error checking, and there are plenty of adaptations to make to this to improve it.
 This is fairly nasty code - but it is a first iteration to get some of the job done.
 This code is currently tailored to printing proposals.
 The final Prioritise.docx, which will be in the same folder as this script, can be opened and printed two to a page.
 Print it by going in print printer propeties  -> Priniting Shortcuts -> Pages per sheet = 2 pages per sheet
"""


import docx
import requests
from six import moves

# Globals - used in a couple of locations


BASE_URL_FOR_ISSUE = r"https://api.github.com/repos/ISISComputingGroup/IBEX/issues"
BASE_URL_PROPOSALS = r"https://api.github.com/search/issues?q=repo:ISISComputingGroup/" \
                     r"IBEX+type:issue+state:open+label:{label}&page={page_number}&per_page={num_per_page}"


def get_web_content_as_json(url, auth):
    """
    This will go and get the contents listed at a URL, and bring it back as a JSON object.
    This needs a try...catch implemented

    Args:
        url: url to get json from
        auth: authentication

    Returns: contents of url as a dictionary

    """
    page_to_get = requests.get(url, auth=auth)
    contents = page_to_get.json()
    try:
        message = contents["message"]
        raise Exception(f"URL returned message: {message}")
    except (KeyError, TypeError):
        pass
    page_to_get.close()
    return contents


class IbexTicket:
    """
    Give our use of labels, this is a class to deal with just those aspects needed for printing for prioritisation.
    """

    # Internal constants
    LabelsToPrioritise = ["proposal", "ready", "in progress"]

    # Fields
    PrioritiseMe = False
    Number = 0000
    Title = "Default Ticket Title"
    Author = "An IBEX Team Member"
    Labels = []
    Proposer = "An IBEX Team Member"
    Estimate = "Un-estimated"

    # Initialisation, the two modes were towards development when some limitation was needed to avoid hitting the
    # API service too much
    def __init__(self, entry, auth, from_file=False):
        if from_file:
            # The file format is known, and so can be treated in a hack-y way
            data = entry.split(";")
            self.Number = int(data[0][1:])
            self.Title = data[1]
            self.Author = data[2]
            labels = data[3]
            labels = labels.replace("[", "")
            labels = labels.replace("]", "")
            labels = labels.replace("'", "")
            labels = labels.replace(",", "")
            self.Labels = labels.split()
            self.Proposer = data[4][:-1]
            self.priority_decision(from_file)
            # See if there is already an estimate for the ticket
            self.get_estimate()
        else:
            # This will be JSON information that can be interpreted
            # Error checking should be added here
            self.Number = entry["number"]
            self.Title = entry["title"]
            self.Author = entry["user"]["login"]
            # Labels need to be dealt with in a slightly different manner as these are a nested JSON structure
            remote_labels = entry["labels"]
            labels = []
            for label in remote_labels:
                labels.append(str(label["name"]))
            self.Labels = labels
            # If there are any labels that need prioritising, make sure they are listed for prioritisation
            if any(ltp in self.LabelsToPrioritise for ltp in labels):
                self.priority_decision()
            # See if there is already an estimate for the ticket
            self.get_estimate()

    def get_estimate(self):
        # Look for an estimate in the labels and apply a value if appropriate
        for labels in self.Labels:
            if is_number(labels):
                self.Estimate = float(labels)
                self.Labels.remove(labels)

    def print_ticket(self):
        # So that you can read it in a console window
        print("Number: %d" % self.Number)
        print("Title: %s" % self.Title)
        print("Author: %s" % self.Author)
        print("Labels: %s" % self.Labels)
        print("Proposer: %s" % self.Proposer)
        print("Estimate: %s" % self.Estimate)

    def file_line(self):
        # Print the bare minimum to the file so that development on the data can continue
        return "#%d;%s;%s;%s;%s\r" % (self.Number, self.Title, self.Author, self.Labels, self.Proposer)

    def priority_decision(self, from_file=False):
        # Decide if this should be prioritised
        # If it is ready, we ought to consider it
        if not self.Labels:
            return
        if "ready" in self.Labels:
            # If it is rework we don't want to consider it, so if it ready, rework should not be an associated label
            if "rework" not in self.Labels:
                self.Proposer = ""
                self.PrioritiseMe = True
        else:
            self.PrioritiseMe = True
            if from_file:
                self.Proposer = ""
            else:
                # Get the events to find the most recent proposer
                url = BASE_URL_FOR_ISSUE + "/" + str(self.Number) + "/events"

                contents = get_web_content_as_json(url, auth)

                proposers = {}
                # Get the list of proposers from the events
                for event in contents:
                    if event["event"] == "labeled":
                        if event["label"]["name"] == "proposal":
                            proposers[event["created_at"]] = event["actor"]["login"]
                # If there is only one proposer, use that proposer ID, otehrwise sort the proposal creation times,
                # and use the proposer from the last one in the list
                if proposers:
                    times = sorted(proposers.keys())
                    self.Proposer = proposers[times[-1]]


def get_open_tickets(auth, label):
    """
    Get all the open tickets
    Will have to page the readbacks
    Note that page 0 and 1 are the same page!
    A Page that is too big comes back with no data, just []
    Args:
        auth: autherentication

    Returns: tickets
    """
    tickets = []

    url = BASE_URL_PROPOSALS.format(page_number=1, num_per_page=1, label=label)

    contents = get_web_content_as_json(url, auth)
    total_open = contents["total_count"]
    # Only need to open enough pages to view the open ones, rather than check for an empty page,
    # calculate the number of pages needed.
    # An error check should be considered for the next iteration.
    number_of_pages = (total_open/100) + 1

    # Open the file for writing during paignation.
    with open("out.txt", "w") as f:
        # Open each page, looking at 100 open issues
        for page_number in range(number_of_pages):
            print(f"Looking at page {page_number + 1}")
            url = BASE_URL_PROPOSALS.format(page_number=page_number + 1, num_per_page=100, label=label)
            contents = get_web_content_as_json(url, auth)["items"]
            for entry in contents:
                # Use a local variable as using in two places at this point
                dev = IbexTicket(entry, auth)
                # Add to the Tickets global for use in other parts of the program
                # Next iteration shuld just return this list
                tickets.append(dev)
                # Write to the file as you go
                f.write(dev.file_line())

    return tickets


def load_tickets_from_file(filename):
    """
    Load tickets from file
    Args:
        filename:

    Returns:

    """
    tickets = []
    # Load the file created to aid in development
    with open(filename, "r") as f:
        for line in f:
            tickets.append(IbexTicket(line, auth=None, from_file=True))
    return tickets


def is_number(number):
    """
    Check for a number, and swallow the error
    Args:
        number: number as string

    Returns:True if it is a number; False otherwise

    """
    try:
        float(number)
        return True
    except ValueError:
        return False


def add_info(document, printable, alignment=docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT, point=40, bold=False,
                 italic=False):
    # Add the details to the document for easy printing
    Paragraph = document.add_paragraph()
    Paragraph.alignment = alignment
    Run = Paragraph.add_run(printable)
    Font = Run.font
    Font.size = docx.shared.Pt(point)
    Run.bold = bold
    Run.italic = italic


###########
##########
#########
# Above here be classes, below be stuff to run
#########
##########
###########

def get_tickets_generate_doc(auth, from_file, label):
    """
    Get the tickets and generate the document in word
    Args:
        auth: wuthentication to use for github
        from_file: True to get data from a file; False otherwise

    """
    # Use one or the other of these lines, depending on what you are allowed to do, and what you have available
    if from_file:
        tickets = load_tickets_from_file("out.txt")
    else:
        tickets = get_open_tickets(auth, label)

    # Just in case you want to look at it
    # for Ticket in Tickets:
    #    Ticket.print_ticket()

    # Create a document for printing - this needs to be moved into a function/class
    document = docx.Document()

    # Initialise the document, change the orientation
    section = document.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    # Loop through the tickets to add the information to the page ready for printing
    # Should consider adding the other labels as well...
    for Ticket in tickets:
        if Ticket.PrioritiseMe:
            add_info(document, str(Ticket.Estimate), alignment=docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT, point=36)
            num_size = 96
            title_size = 48
            if len(Ticket.Title) > 35:
                num_size = 84
                title_size = 32
            add_info(document, "#" + str(Ticket.Number), point=num_size, bold=True)
            add_info(document, Ticket.Title, point=title_size)
            number_of_rows = max(2, len(Ticket.Labels))
            table = document.add_table(rows=number_of_rows, cols=2)
            people_cells = table.columns[0].cells
            add_info(people_cells[0], "Author: " + Ticket.Author, point=20)
            add_info(people_cells[1], "Proposed by: " + Ticket.Proposer, italic=True, point=14)
            label_cells = table.columns[1].cells
            if Ticket.Labels:
                label_index = 0
                for label in Ticket.Labels:
                    if not is_number(label):
                        add_info(label_cells[label_index], label, point=14)
                        label_index = label_index + 1
            document.add_page_break()

    # Save the document so that you can load and print it
    document.save("Prioritise.docx")


if __name__ == "__main__":
    username = moves.input("Username (blank for no auth): ")
    if username != "":
        token = moves.input("Token (generated in personal tokens): ")
        auth_pair = (username, token)
    else:
        auth_pair = None

    from_file = moves.input("From file (Y or N): ").upper()[0] == "Y"
    label = moves.input("Label (p - proposal, r - ready, i - in+progress, anything else): ")
    if label[0] == "p":
        label = "proposal"
    elif label[0] == "r":
        label = "ready"
    elif label[0] == "i":
        label = '"in+progress"'

    print(label)
    get_tickets_generate_doc(auth_pair, from_file, label)
