# A Utility to check and provide a printable list of proposed and ready tickets ready for prioritisation.
# To use this the python-docx package needs to be installed.
# There is limited error checking, and there are plenty of adaptations to make to this to improve it.
# This is fairly nasty code - but it is a first iteration to get some of the job done.
# This code is currently tailored to printing proposals.
# The final Prioritise.docx, which will be in the same folder as this script, can be opened and printed two to a page.

# The required imports
import urllib2
import json
import docx

# Globals - used in a couple of locations
BaseURL = "https://api.github.com/repos/ISISComputingGroup/IBEX/issues"
Tickets = []


def get_web_content_as_json(url):
    # This will go and get the contents listed at a URL, and bring it back as a JSON object.
    # This needs a try...catch implemented
    page_to_get = urllib2.urlopen(url)
    contents = json.loads(page_to_get.read())
    page_to_get.close()

    return contents


class IbexTicket:
    # Give our use of labels, this is a class to deal with just those aspects needed for printing for prioritisation.

    # Internal constants
    LabelsToPrioritise = ["proposal", "ready"]

    # Fields
    PrioritiseMe = False
    Number = 0000
    Title = "Default Ticket Title"
    Author = "An IBEX Team Member"
    Labels = []
    Proposer = "An IBEX Team Member"
    Estimate = "Un-estimated"

    # Initialisation, the two modes were towards development when some limitation was needed to avoid hitting the
    # API service too much
    def __init__(self, entry, from_file=False):
        if from_file:
            # The file format is known, and so can be treated in a hack-y way
            data = entry.split(";")
            self.Number = int(data[0][1:])
            self.Title = data[1]
            self.Author = data[2]
            labels = data[3]
            labels = labels.replace("[", "")
            labels = labels.replace("]", "")
            labels = labels.replace("'", "")
            labels = labels.replace(",", "")
            self.Labels = labels.split()
            self.Proposer = data[4][:-1]
            self.priority_decision(from_file)
            # See if there is already an estimate for the ticket
            self.get_estimate()
        else:
            # This will be JSON information that can be interpreted
            # Error checking should be added here
            self.Number = entry["number"]
            self.Title = entry["title"]
            self.Author = entry["user"]["login"]
            # Labels need to be dealt with in a slightly different manner as these are a nested JSON structure
            remote_labels = entry["labels"]
            labels = []
            for label in remote_labels:
                labels.append(str(label["name"]))
            self.Labels = labels
            # If there are any labels that need prioritising, make sure they are listed for prioritisation
            if any(ltp in self.LabelsToPrioritise for ltp in labels):
                self.priority_decision()
            # See if there is already an estimate for the ticket
            self.get_estimate()

    def get_estimate(self):
        # Look for an estimate in the labels and apply a value if appropriate
        for labels in self.Labels:
            if is_number(labels):
                self.Estimate = float(labels)
                self.Labels.remove(labels)

    def print_ticket(self):
        # So that you can read it in a console window
        print("Number: %d" % self.Number)
        print("Title: %s" % self.Title)
        print("Author: %s" % self.Author)
        print("Labels: %s" % self.Labels)
        print("Proposer: %s" % self.Proposer)
        print("Estimate: %s" % self.Estimate)

    def file_line(self):
        # Print the bare minimum to the file so that development on the data can continue
        return "#%d;%s;%s;%s;%s\r" % (self.Number, self.Title, self.Author, self.Labels, self.Proposer)

    def priority_decision(self, from_file=False):
        # Decide if this should be prioritised
        # If it is ready, we ought to consider it
        if not self.Labels:
            return
        if "ready" in self.Labels:
            # If it is rework we don't want to consider it, so if it ready, rework should not be an associated label
            if "rework" not in self.Labels:
                self.Proposer = ""
                self.PrioritiseMe = True
        else:
            self.PrioritiseMe = True
            if from_file:
                self.Proposer = ""
            else:
                # Get the events to find the most recent proposer
                url = BaseURL + "/" + str(self.Number) + "/events"
                contents=get_web_content_as_json(url)
                proposers = {}
                # Get the list of proposers from the events
                for event in contents:
                    if event["event"] == "labeled":
                        if event["label"]["name"] == "proposal":
                            proposers[event["created_at"]] = event["actor"]["login"]
                # If there is only one proposer, use that proposer ID, otehrwise sort the proposal creation times,
                # and use the proposer from the last one in the list
                if proposers:
                    times = sorted(proposers.keys())
                    self.Proposer = proposers[times[-1]]


def get_open_tickets():
    # Get all the open tickets
    # Will have to page the readbacks
    # Note that page 0 and 1 are the same page!
    # A Page that is too big comes back with no data, just []

    url = "https://api.github.com/search/issues?q=repo:ISISComputingGroup/IBEX+type:issue+state:open"
    contents = get_web_content_as_json(url)
    total_open = contents["total_count"]
    # Only need to open enough pages to view the open ones, rather than check for an empty page,
    # calculate the number of pages needed.
    # An error check should be considered for the next iteration.
    number_of_pages = (total_open/100)+1

    # Open the file for writing during paignation.
    f = open("out.txt", "w")

    # Open each page, looking at 100 open issues
    for pageNumber in range(1,number_of_pages,1):
        print("Looking at page %d" % pageNumber)
        url = BaseURL + "?state=open&page=" + str(pageNumber) + "&per_page=100"
        contents = get_web_content_as_json(url)

        for entry in contents:
            # Use a local variable as using in two places at this point
            dev = IbexTicket(entry)
            # Add to the Tickets global for use in other parts of the program
            # Next iteration shuld just return this list
            Tickets.append(dev)
            # Write to the file as you go
            f.write(dev.file_line())

    # Close the file to be tidy.
    f.close()


def load_text(filename):
    # Load the file created to aid in development
    f = open(filename,"r")
    for line in f:
        Tickets.append(IbexTicket(line, from_file=True))
    f.close()


def is_number(s):
    # Check for a number, and swallow the error
    try:
        float(s)
        return True
    except ValueError:
        return False


def add_info(document, printable, alignment=docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT, point=40, bold=False,
                 italic=False):
    # Add the details to the document for easy printing
    Paragraph = document.add_paragraph()
    Paragraph.alignment = alignment
    Run = Paragraph.add_run(printable)
    Font = Run.font
    Font.size = docx.shared.Pt(point)
    Run.bold = bold
    Run.italic = italic


###########
##########
#########
# Above here be classes, below be stuff to run
#########
##########
###########

# Use one or the other of these lines, depending on what you are allowed to do, and what you have available
# loadText("out.txt")
# get_open_tickets()
load_text("toformat.txt")

# Just in case you want to look at it
# for Ticket in Tickets:
#    Ticket.print_ticket()

# Create a document for printing - this needs to be moved into a function/class
document = docx.Document()

# Initialise the document, change the orientation
section = document.sections[0]
new_width, new_height = section.page_height, section.page_width
section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height

# Loop through the tickets to add the information to the page ready for printing
# Should consider adding the other labels as well...
for Ticket in Tickets:
    if Ticket.PrioritiseMe:
        add_info(document,str(Ticket.Estimate),alignment=docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT, point=36)
        num_size = 96
        title_size = 48
        if len(Ticket.Title) > 35:
            num_size = 84
            title_size = 32
        add_info(document, "#" + str(Ticket.Number), point=num_size, bold=True)
        add_info(document, Ticket.Title, point=title_size)
        number_of_rows = max(2,len(Ticket.Labels))
        table = document.add_table(rows=number_of_rows, cols=2)
        people_cells = table.columns[0].cells
        add_info(people_cells[0],"Author: " + Ticket.Author, point=20)
        add_info(people_cells[1],"Proposed by: " + Ticket.Proposer, italic=True, point=14)
        label_cells = table.columns[1].cells
        if Ticket.Labels:
            label_index = 0
            for label in Ticket.Labels:
                if not is_number(label):
                    add_info(label_cells[label_index],label, point=14)
                    label_index = label_index + 1
        document.add_page_break()

# Save the document so that you can load and print it
document.save("Prioritise.docx")


